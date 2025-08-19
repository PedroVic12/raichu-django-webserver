from docx import Document
from docx.shared import Inches
from typing import List, Tuple


class SalamanceWord:
    def __init__(self, file_name: str):
        self.document = Document()
        self.file_name = file_name

    def add_heading(self, text: str, level: int):
        self.document.add_heading(text, level)

    def add_paragraph(self, text: str, style: str = None):
        self.document.add_paragraph(text, style=style)

    def add_picture(self, image_path: str, width: float):
        self.document.add_picture(image_path, width=Inches(width))

    def add_table(self, rows: int, cols: int, headers: List[str] = None, records: List[Tuple] = None):
        table = self.document.add_table(rows=rows, cols=cols)
        if headers:
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
        if records:
            for record in records:
                row_cells = table.add_row().cells
                for i, item in enumerate(record):
                    row_cells[i].text = str(item)

    def add_page_break(self):
        self.document.add_page_break()

    def save(self):
        self.document.save(self.file_name)


class DemoDocBuilder:
    def __init__(self, file_name: str):
        self.builder = SalamanceWord(file_name)

    def build(self):
        self.builder.add_heading('Document Title', 0)

        self.builder.add_paragraph('A plain paragraph having some ', style=None)
        self.builder.add_paragraph('bold', style=None)
        self.builder.add_paragraph(' and some ', style=None)
        self.builder.add_paragraph('italic.', style=None)

        self.builder.add_heading('Heading, level 1', level=1)
        self.builder.add_paragraph('Intense quote', style='Intense Quote')

        self.builder.add_paragraph('first item in unordered list', style='List Bullet')
        self.builder.add_paragraph('first item in ordered list', style='List Number')

        self.builder.add_picture(r'C:\Users\PedroVictorRodrigues\Documents\GitHub\elon-musk\Tecnologia e Inovação\Visão Computacional\assets\OS Eletrica\OCR\ocr_relatorioOS (3).jpg', width=1.25)

        records = (
            (3, '101', 'Spam'),
            (7, '422', 'Eggs'),
            (4, '631', 'Spam, spam, eggs, and spam')
        )

        self.builder.add_table(rows=1, cols=3, headers=['Qty', 'Id', 'Desc'], records=records)

        self.builder.add_page_break()
        self.builder.save()


if __name__ == "__main__":
    demo_builder = DemoDocBuilder('demo.docx')
    demo_builder.build()
